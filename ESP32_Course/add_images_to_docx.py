import os
import docx
from docx.shared import Inches

def get_mapping():
    return {
        "NTC_Thermistor_10k.jpg": "ntc thermistor",
        "PT100_MAX31865.jpg": "pt100",
        "TMP36.jpg": "tmp36",
        "DS18B20.jpg": "ds18b20",
        "SHT31.jpg": "sht31",
        "MAX6675_thermocouple.jpg": "max6675",
        "MLX90614_infrared_temperature_sensor.jpg": "mlx90614",
        "SHTC3_humidity_sensor.jpg": "shtc3",
        "DHT22_humidity_sensor.jpg": "dht22",
        "BME280_sensor_module.jpg": "bme280",
        "MQ-135_gas_sensor_module.jpg": "mq-135",
        "MH-Z19B_CO2_sensor.jpg": "mh-z19b",
        "SGP30_VOC_sensor.jpg": "sgp30",
        "BME680_gas_sensor.jpg": "bme680",
        "PMS5003_dust_sensor.jpg": "pms5003",
        "GL5528_LDR_sensor.jpg": "gl5528",
        "BH1750_light_sensor.jpg": "bh1750",
        "TSL2561_light_sensor.jpg": "tsl2561",
        "TCS34725_color_sensor.jpg": "tcs3472",
        "APDS9999_sensor_module.jpg": "apds9999",
        "VEML6075_UV_sensor.jpg": "veml6075",
        "AS7262_spectral_sensor.jpg": "as7262",
        "Capacitive_soil_moisture_sensor_v2.0.jpg": "capacitive",
        "DFRobot_EC_sensor.jpg": "dfrobot",
        "Gravity_pH_sensor.jpg": "gravity ph",
        "MPU6050_module.jpg": "mpu6050",
        "MPU9250_module.jpg": "mpu9250",
        "ADXL345_accelerometer.jpg": "adxl345",
        "HC-SR04_ultrasonic_sensor.jpg": "hc-sr04",
        "VL53L0X_ToF_sensor.jpg": "vl53l0x",
        "GP2Y0A21_IR_distance_sensor.jpg": "gp2y0a21",
        "BMP280_pressure_sensor.jpg": "bmp280",
        "HX711_load_cell_amplifier.jpg": "hx711",
        "FSR402_force_sensor.jpg": "fsr402",
        "A3144_hall_sensor.jpg": "a3144",
        "ACS712_current_sensor.jpg": "acs712",
        "INA219_current_sensor.jpg": "ina219",
        "MAX17048_fuel_gauge.jpg": "max17048",
        "OV2640_camera_module.jpg": "ov2640",
        "FLIR_Lepton_thermal_camera.jpg": "flir",
        "Analog_microphone_module.jpg": "아날로그 마이크",
        "INMP441_MEMS_microphone.jpg": "inmp441",
        "TTP223_touch_sensor.jpg": "ttp223",
        "ESP32_capacitive_touch.jpg": "esp32 터치핀",
        "YF-S201_water_flow_sensor.jpg": "yf-s201",
        "Float_switch_water_level_sensor.jpg": "float switch",
        "Turbidity_sensor_module.jpg": "turbidity"
    }

from docx.enum.text import WD_LINE_SPACING

def main():
    doc_path = 'ESP32_Sensor_Catalog_KR_clean.docx'
    doc = docx.Document(doc_path)
    mapping = get_mapping()
    
    inserted = set()
    for para in doc.paragraphs:
        txt = para.text.lower()
        if not txt.strip():
            continue
            
        for img_file, kw in mapping.items():
            if kw in txt and img_file not in inserted:
                img_path = os.path.join('sensor_images', img_file)
                if os.path.exists(img_path):
                    try:
                        new_p = para.insert_paragraph_before('')
                        new_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        run = new_p.add_run()
                        run.add_picture(img_path, width=Inches(1.5))
                        inserted.add(img_file)
                        print(f"Inserted {img_file} into DOCX!")
                    except Exception as e:
                        print(f"Error adding picture {img_file}: {e}")
                        
    doc.save('ESP32_Sensor_Catalog_KR.docx')
    print(f"Done adding {len(inserted)} images to DOCX.")

if __name__ == '__main__':
    main()
